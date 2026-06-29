import os
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PLOTS_DIR = os.path.join(BASE_DIR, "plots")
ARTIFACT_DIR = r"C:\Users\Ria\.gemini\antigravity\brain\5b70e094-feec-4d9f-8983-1b0aa0aa9b7c"

# Colors for styling
DARK_BLUE = "1B365D"
LIGHT_BLUE = "F0F4F8"
CHARCOAL = "333333"
WHITE = "FFFFFF"

def set_cell_background(cell, color_hex):
    """Sets background color of a cell."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell margins (padding) in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_paragraph(doc, text, style_name="Normal", space_before=0, space_after=6, bold=False, italic=False, color=None, font_size=11):
    """Creates a paragraph with explicit spacing and font settings."""
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def create_report():
    doc = Document()
    
    # Page setup - Margins (1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Set default style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor.from_string(CHARCOAL)
    
    # ----------------------------------------------------
    # COVER PAGE (Modern Minimalist Style)
    # ----------------------------------------------------
    # Title
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(100)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = p_title.add_run("BUILDING A TRUSTWORTHY\nBACKTESTING PIPELINE")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(28)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p_title.paragraph_format.space_after = Pt(10)
    
    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_sub = p_sub.add_run("Task 2 Research Report: Infrastructure, Baselines, and Research Proposals")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(14)
    run_sub.italic = True
    run_sub.font.color.rgb = RGBColor.from_string("666666")
    p_sub.paragraph_format.space_after = Pt(150)
    
    # Metadata Block
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(4)
    r_meta1 = p_meta.add_run("Assignee: ")
    r_meta1.bold = True
    p_meta.add_run("Ria Chawak\n")
    r_meta2 = p_meta.add_run("Affiliation: ")
    r_meta2.bold = True
    p_meta.add_run("Research Intern, Indian Institute of Technology Bombay (IITB)\n")
    r_meta3 = p_meta.add_run("Mentor: ")
    r_meta3.bold = True
    p_meta.add_run("Project Mentor / Professor\n")
    r_meta4 = p_meta.add_run("Date: ")
    r_meta4.bold = True
    p_meta.add_run("June 2026\n")
    
    doc.add_page_break()
    
    # ----------------------------------------------------
    # SECTION 1: EXECUTIVE SUMMARY
    # ----------------------------------------------------
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("1. Executive Summary")
    r.font.size = Pt(18)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    
    add_styled_paragraph(doc, 
        "This research report documents the engineering workflow and mathematical validation of a custom backtesting simulator, "
        "fulfilling the requirements of Task 2 for the quantitative research internship at IIT Bombay. Rigorous quantitative trading "
        "requires infrastructure that is immune to logical errors, specifically look-ahead bias, and fully cost-aware. We establish "
        "a reliable backtesting pipeline that processes daily (EOD) historical price data for both Crypto Spot and Indian Equities. "
        "We implement three classical baseline trading strategies—Trend Following (SMA Crossover), Mean Reversion (Bollinger Bands), "
        "and Statistical Arbitrage (Spread Z-score Reversion)—and benchmark them on our custom simulator and the industry-standard "
        "backtesting.py package. Finally, we provide a focused literature review on deep learning/machine learning applications in "
        "financial markets and propose three concrete, falsifiable research tracks to beat the classical baselines established here."
    )
    
    # ----------------------------------------------------
    # SECTION 2: MARKET AND FREQUENCY SELECTION
    # ----------------------------------------------------
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    r = h2.add_run("2. Market & Frequency Selection")
    r.font.size = Pt(18)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    
    add_styled_paragraph(doc, 
        "A critical phase of quantitative research is selecting the asset universe and trade frequency. For this research foundation, "
        "we make a deliberate choice to support two distinct markets representing different regulatory, liquidity, and structural features:"
    )
    
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Crypto Spot (BTCUSDT & ETHUSDT): ").bold = True
    p.add_run("Crypto spot markets offer a 24/7 continuous trading environment, which eliminates the complexities of overnight price gaps "
              "and market holiday anomalies. This serves as a clean starting point for statistical modeling. Due to its high retail participation "
              "and lack of circuit breakers, crypto assets exhibit high volatility and strong trend-following traits, making them ideal "
              "testbeds for algorithmic trading strategies.")
              
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(6)
    p.add_run("Indian Equities (TCS & INFY): ").bold = True
    p.add_run("To ensure our pipeline is robust under real-world regulatory restrictions and high trading frictions, we test on two large-cap "
              "constituents of the NSE. Trading in the Indian market requires meticulous modeling of specific frictions, including GST, "
              "Securities Transaction Tax (STT), exchange transaction charges, stamp duties, and SEBI turnover fees. Modeling these frictions "
              "ensures that backtest results are realistic and prevents the inflation of equity curves.")
              
    add_styled_paragraph(doc, 
        "Frequency Selection: Daily (EOD) data is selected as the base frequency. Starting with daily data allows us to prove the stability "
        "of our backtest simulator end-to-end without the infrastructural burden, database size, and latency issues associated with intraday or "
        "Limit Order Book (LOB) tick data. The models can be scaled down to lower timeframes (e.g., hourly or 15-minute) once the rig is validated."
    )
    
    # ----------------------------------------------------
    # SECTION 3: EXPLORATORY DATA ANALYSIS (EDA)
    # ----------------------------------------------------
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(18)
    h3.paragraph_format.space_after = Pt(6)
    r = h3.add_run("3. Exploratory Data Analysis")
    r.font.size = Pt(18)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    
    add_styled_paragraph(doc, 
        "A reproducibility check was run on the downloaded daily data spanning from January 1, 2018, to June 2026. The statistical metrics "
        "of the daily returns are summarized in Table 1 below:"
    )
    
    # Table 1: EDA Summary
    table_eda = doc.add_table(rows=5, cols=6)
    table_eda.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers_eda = ["Asset", "Bars", "Mean Return", "Ann. Volatility", "Skewness", "Excess Kurtosis"]
    col_widths_eda = [Inches(1.2), Inches(0.8), Inches(1.1), Inches(1.2), Inches(0.9), Inches(1.2)]
    
    # Format Header
    for idx, name in enumerate(headers_eda):
        cell = table_eda.cell(0, idx)
        cell.text = name
        set_cell_background(cell, DARK_BLUE)
        set_cell_margins(cell)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    data_eda = [
        ["BTCUSDT", "3,097", "0.001076", "64.90%", "-0.3702", "9.2524"],
        ["ETHUSDT", "3,097", "0.001261", "85.41%", "-0.1793", "6.6879"],
        ["TCS", "2,095", "0.000443", "24.41%", "0.0007", "4.3400"],
        ["INFY", "2,095", "0.000607", "27.49%", "-0.3573", "8.4234"]
    ]
    
    for r_idx, row_data in enumerate(data_eda):
        row = table_eda.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            set_cell_margins(cell)
            if r_idx % 2 == 1:
                set_cell_background(cell, LIGHT_BLUE)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    # Set column widths
    for row in table_eda.rows:
        for idx, width in enumerate(col_widths_eda):
            row.cells[idx].width = width
            
    add_styled_paragraph(doc, 
        "Key Data Quirks and Statistical Observations:\n"
        "1. Heavy Tails (Excess Kurtosis): All four assets exhibit high excess kurtosis (values significantly greater than 0). "
        "BTCUSDT (9.25) and INFY (8.42) have extremely fat tails (leptokurtic), implying that extreme returns (market crashes or vertical spikes) "
        "occur far more frequently than a standard normal distribution models. Any ML model must account for these non-Gaussian distributions.\n"
        "2. Skewness: The negative skewness of BTC, ETH, and INFY indicates a left-skewed distribution, showing that downward moves are often "
        "more rapid and severe than upward moves (volatility clustering during panics). TCS exhibits almost symmetric returns (skewness ~0.0007).\n"
        "3. Continuity: Crypto data is 100% gapless across the 3,097 bars. For Indian equities, the data-loading engine successfully filtered "
        "weekend and holiday gaps, identifying 49 long weekends where market inactivity spanned more than 3 calendar days (normal NSE schedule)."
    )
    
    # ----------------------------------------------------
    # SECTION 4: BACKTESTER ENGINE DESIGN & VALIDATION
    # ----------------------------------------------------
    h4 = doc.add_paragraph()
    h4.paragraph_format.space_before = Pt(18)
    h4.paragraph_format.space_after = Pt(6)
    r = h4.add_run("4. Custom Backtester Engine Design & Validation")
    r.font.size = Pt(18)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    
    add_styled_paragraph(doc, 
        "The core deliverable of this task is a custom backtesting engine written from scratch. "
        "In quant research, the backtesting engine is the foundation; a leaky engine renders any ML/DL model useless. "
        "Our custom engine is designed around two strict principles: No Look-Ahead Bias and Cost/Slippage Awareness."
    )
    
    add_styled_paragraph(doc, "4.1 No Look-Ahead Bias Execution Model", bold=True)
    add_styled_paragraph(doc, 
        "To prevent look-ahead bias (using future information to trade in the present), our engine implements a strict 1-bar execution delay. "
        "A trading decision (signal) generated at the close of bar t (using Close price P_t, or rolling indicators computed up to day t) "
        "is executed at the Open price of day t+1 (P_open,t+1). This is mathematically represented as:"
    )
    
    add_styled_paragraph(doc, 
        "Trade Execution Price = P_open,t+1 * (1 +/- lambda_slippage)\n"
        "Target Portfolio Position = Weight_t * Equity_t+1,open", italic=True
    )
    
    add_styled_paragraph(doc, "4.2 Transaction Fee & Slippage Models", bold=True)
    add_styled_paragraph(doc, 
        "Frictions can destroy a strategy's profitability. Our engine supports two distinct cost structures:"
    )
    
    p_eq = doc.add_paragraph()
    p_eq.paragraph_format.left_indent = Inches(0.5)
    p_eq.add_run("1. Crypto Cost Model: ").bold = True
    p_eq.add_run("A flat commission of 0.1% (representing standard Binance Taker commissions) plus 0.05% slippage, representing a flat 0.15% friction per order value.\n")
    p_eq.add_run("2. Indian Equities Cost Model (Delivery-based): ").bold = True
    p_eq.add_run("Frictions are calculated dynamically based on NSE delivery rules:\n"
                 "- Brokerage = min(0.0003 * Trade_Value, 20.0 INR)\n"
                 "- STT (Securities Transaction Tax) = 0.1% * Trade_Value (Buy and Sell)\n"
                 "- Exchange Transaction Charges = 0.00345% * Trade_Value\n"
                 "- SEBI Turnover Fee = 0.0001% * Trade_Value\n"
                 "- Stamp Duty = 0.015% * Trade_Value (Buy orders only)\n"
                 "- GST = 18% * (Brokerage + Exchange Charges + SEBI Fee)\n"
                 "- Slippage = 0.05% * Trade_Value")
                 
    add_styled_paragraph(doc, "4.3 Engine Sanity Check Validation", bold=True)
    add_styled_paragraph(doc, 
        "To validate the simulator, we executed two sanity tests before deploying our baseline strategies:\n"
        "1. Buy & Hold Validation: If we set target weights to 1.0 (fully invested) for all bars, the gross equity curve return must match the "
        "asset's return from the first execution (Day 1 Open) to the final Close. Our engine passed this test, with the gross B&H return "
        "matching the asset's Open-to-Close return exactly (within a 1e-4 tolerance, returning a successful match of -9.65% gross vs -9.65% asset return on the dummy validation set).\n"
        "2. Random Signal Validation: We generated random signals (alternating buy/sell every 20 bars). Over a long timeframe, a zero-edge strategy "
        "must underperform the gross curve due to transaction fees. The engine confirmed a clear cost drag (Gross return = -5.57%, Net return = -5.86%, "
        "representing a cost drag of 0.28%), validating that frictions are correctly accumulated on each transaction."
    )
    
    # ----------------------------------------------------
    # SECTION 5: CLASSICAL BASELINE STRATEGIES
    # ----------------------------------------------------
    h5 = doc.add_paragraph()
    h5.paragraph_format.space_before = Pt(18)
    h5.paragraph_format.space_after = Pt(6)
    r = h5.add_run("5. Classical Baseline Strategies & Results")
    r.font.size = Pt(18)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    
    add_styled_paragraph(doc, 
        "We implement and evaluate three classical strategies on both the Crypto Spot (BTCUSDT) and Indian Equities (TCS & INFY) markets. "
        "These baselines serve as the benchmarks that any subsequent machine learning models must beat."
    )
    
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Strategy 1: SMA Crossover (Trend Following): ").bold = True
    p.add_run("Uses the 50-day and 200-day Simple Moving Averages. A long signal (1.0) is triggered when the 50-day SMA crosses above the 200-day SMA. "
              "A cash exit (0.0) is triggered when it crosses below. For crypto, it captures major bull markets.")
              
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Strategy 2: Bollinger Band Mean Reversion: ").bold = True
    p.add_run("Calculates the 20-day SMA and +/- 2 standard deviation bands. When price crosses below the lower band (oversold), we go long. "
              "We close the position when price reverts to the 20-day SMA. For crypto, we also test a Long/Short version (shorting the upper band).")
              
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(6)
    p.add_run("Strategy 3: Statistical Arbitrage / Pairs Trading: ").bold = True
    p.add_run("Tests cointegration between BTCUSDT & ETHUSDT, and TCS & INFY. Using a 60-day rolling lookback, we fit a linear regression spread. "
              "When the z-score of the spread exceeds 1.5, we short the expensive asset and buy the cheap asset (50% weight each). "
              "We exit when the z-score reverts to 0.")
              
    add_styled_paragraph(doc, 
        "The quantitative metrics of these strategies are summarized in Table 2:"
    )
    
    # Table 2: Strategy Metrics
    table_strat = doc.add_table(rows=7, cols=6)
    table_strat.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers_strat = ["Asset", "Strategy", "CAGR (%)", "Volatility (%)", "Sharpe", "Max Drawdown (%)"]
    col_widths_strat = [Inches(1.2), Inches(2.2), Inches(0.9), Inches(1.1), Inches(0.8), Inches(1.3)]
    
    for idx, name in enumerate(headers_strat):
        cell = table_strat.cell(0, idx)
        cell.text = name
        set_cell_background(cell, DARK_BLUE)
        set_cell_margins(cell)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    data_strat = [
        ["BTCUSDT", "SMA Crossover (Long-Only)", "26.19%", "46.84%", "0.736", "-66.81%"],
        ["TCS", "SMA Crossover (Long-Only)", "4.82%", "18.32%", "0.354", "-28.53%"],
        ["BTCUSDT", "Bollinger Bands Mean Reversion", "-29.42%", "50.87%", "-0.422", "-96.98%"],
        ["TCS", "Bollinger Bands (Long-Only)", "-3.50%", "12.77%", "-0.221", "-34.77%"],
        ["BTC/ETH", "Pairs Trading (L/S)", "-14.47%", "17.25%", "-0.820", "-74.04%"],
        ["TCS/INFY", "Pairs Trading (L/S)", "-0.43%", "4.20%", "-0.083", "-12.49%"]
    ]
    
    for r_idx, row_data in enumerate(data_strat):
        row = table_strat.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            set_cell_margins(cell)
            if r_idx % 2 == 1:
                set_cell_background(cell, LIGHT_BLUE)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    for row in table_strat.rows:
        for idx, width in enumerate(col_widths_strat):
            row.cells[idx].width = width
            
    add_styled_paragraph(doc, 
        "Market Dynamics & Analysis:\n"
        "1. SMA Crossover (Trend Following): This was the only profitable baseline. On BTCUSDT, it delivered a solid 26.19% CAGR, "
        "although with a massive drawdown of -66.81% during major crypto winters. This highlights the typical profile of trend following: "
        "high risk, high return, and long drawdowns. On TCS, the SMA crossover yielded a modest 4.82% CAGR, reflecting low market momentum.\n"
        "2. Bollinger Bands (Mean Reversion): This strategy performed poorly on both assets (BTC: -29.42% CAGR, TCS: -3.50% CAGR). "
        "In highly trending markets like crypto, mean reversion strategies suffer from the 'falling knife' problem: buying early during "
        "crashes, or shorting early during parabolic rallies, leading to extreme drawdowns (BTC MaxDD: -96.98%).\n"
        "3. Pairs Trading (Statistical Arbitrage): TCS/INFY showed relatively low volatility (4.20%) and a minor drawdown (-12.49%), "
        "demonstrating the market-neutral property of statistical arbitrage. However, it was slightly unprofitable (-0.43% CAGR) due to "
        "transaction fee drag on frequent rebalancing. The BTC/ETH pair experienced a severe divergence (MaxDD: -74.04%), proving that "
        "crypto cointegration relationships are highly unstable (non-stationary) over multi-year horizons."
    )
    
    # Insert Charts
    doc.add_page_break()
    h_charts = doc.add_paragraph()
    h_charts.paragraph_format.space_before = Pt(12)
    h_charts.paragraph_format.space_after = Pt(6)
    r_c = h_charts.add_run("Equity Curves and Drawdowns")
    r_c.font.size = Pt(14)
    r_c.bold = True
    r_c.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    
    charts = [
        ("btc_sma_crossover.png", "Figure 1: BTCUSDT SMA Crossover Equity Curve"),
        ("tcs_sma_crossover.png", "Figure 2: TCS SMA Crossover Equity Curve"),
        ("btc_bb_reversion.png", "Figure 3: BTCUSDT Bollinger Bands Reversion Equity Curve"),
        ("tcs_bb_reversion.png", "Figure 4: TCS Bollinger Bands Reversion Equity Curve"),
        ("btc_eth_pairs_trading.png", "Figure 5: BTC/ETH Pairs Trading Equity Curve"),
        ("tcs_infy_pairs_trading.png", "Figure 6: TCS/INFY Pairs Trading Equity Curve")
    ]
    
    for filename, caption in charts:
        path = os.path.join(PLOTS_DIR, filename)
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(5.0))
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_cap = p_cap.add_run(caption)
            run_cap.font.size = Pt(9)
            run_cap.italic = True
            p_cap.paragraph_format.space_after = Pt(18)
            
    doc.add_page_break()
    
    # ----------------------------------------------------
    # SECTION 6: FOCUSED LITERATURE REVIEW
    # ----------------------------------------------------
    h6 = doc.add_paragraph()
    h6.paragraph_format.space_before = Pt(18)
    h6.paragraph_format.space_after = Pt(6)
    r = h6.add_run("6. Focused Literature Review on ML/DL for Trading")
    r.font.size = Pt(18)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    
    add_styled_paragraph(doc, 
        "Machine learning and deep learning models are increasingly used to replace classical rules. However, they introduce significant "
        "risks of overfitting, non-stationarity, and look-ahead bias. Table 3 provides a structured comparison of foundational and recent literature:"
    )
    
    # Table 3: Literature Review
    table_lit = doc.add_table(rows=6, cols=5)
    table_lit.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers_lit = ["Paper", "Task & Model", "Input Data & Market", "Reported Edge", "Limitations & Frictions"]
    col_widths_lit = [Inches(1.5), Inches(1.3), Inches(1.3), Inches(1.4), Inches(1.5)]
    
    for idx, name in enumerate(headers_lit):
        cell = table_lit.cell(0, idx)
        cell.text = name
        set_cell_background(cell, DARK_BLUE)
        set_cell_margins(cell)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    data_lit = [
        [
            "Gu, Kelly & Xiu (2020)\n'Empirical Asset Pricing via ML'",
            "Regression (Trees, Forest, Neural Nets)",
            "US Equities (900+ macro/micro features)",
            "Trees and Neural Nets double the Sharpe ratio of classical linear baselines.",
            "No transaction costs or bid-ask spreads modeled; high turnover poses execution drag."
        ],
        [
            "Krauss, Do & Huck (2017)\n'Deep Nets, GBT, RF for Stat Arb'",
            "Classification (DNN, GBDT, Random Forest)",
            "S&P 500 constituents, Daily data",
            "GBDT/RF outperform DNN, achieving daily returns > 0.40% before costs.",
            "Catastrophic drawdown in 2008-09; returns decay significantly after 2010 once costs are applied."
        ],
        [
            "Fischer & Krauss (2018)\n'LSTMs for Market Prediction'",
            "Sequence Classification (LSTM)",
            "S&P 500 constituents, Daily data (1992-2015)",
            "LSTM outperforms GBDT and Random Forest, capturing time-series dependencies.",
            "Highly sensitive to transaction fees. An EOD fee of 0.15% wipes out most of the LSTM excess returns."
        ],
        [
            "Lim, Zohren & Roberts (2019)\n'Deep Learning for TS Momentum'",
            "Sequence Regression (LSTM + Attention)",
            "50+ Liquid Futures (Commodities, FX, Equities)",
            "Attention-LSTM adaptively weights time-series momentum, avoiding decay in flat markets.",
            "Requires substantial historical data to train attention parameters; execution slippage not modeled."
        ],
        [
            "Gort et al. (2023)\n'Deep RL for Crypto Trading'",
            "Policy Optimization (PPO / Actor-Critic)",
            "Crypto Spot/Derivatives (BTC, ETH, etc.)",
            "Deep RL agents optimize execution entries and adjust size dynamically based on volatility.",
            "Extreme sensitivity to hyperparameter initialization; prone to overfitting to short-term regimes."
        ]
    ]
    
    for r_idx, row_data in enumerate(data_lit):
        row = table_lit.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            set_cell_margins(cell)
            if r_idx % 2 == 1:
                set_cell_background(cell, LIGHT_BLUE)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
    for row in table_lit.rows:
        for idx, width in enumerate(col_widths_lit):
            row.cells[idx].width = width
            
    add_styled_paragraph(doc, 
        "How the Literature Guards Against Overfitting and Leakage:\n"
        "1. Purging and Embargoing (López de Prado): In overlapping labels (e.g., predicting 5-day forward return), adjacent observations "
        "share data, which violates the IID assumption. Purging removes training labels that overlap with testing data, while embargoing "
        "discards observations immediately following the training set to prevent leakage from autoregressive effects.\n"
        "2. Cross-Validation Schemes: Time-series split (walk-forward validation) is preferred over k-fold cross-validation. "
        "Standard k-fold shuffles indices, which allows future data to leak into the past. Walk-forward testing ensures the training set "
        "strictly precedes the validation set.\n"
        "3. Deflated Sharpe Ratio (DSR): Accounts for the number of backtest trials ran. When researchers run hundreds of strategy permutations, "
        "the best-performing strategy is often selected by pure chance. DSR adjusts the Sharpe ratio downward based on the variance of trials "
        "to prevent type-I errors (false positives)."
    )
    
    # ----------------------------------------------------
    # SECTION 7: RESEARCH PROPOSALS
    # ----------------------------------------------------
    h7 = doc.add_paragraph()
    h7.paragraph_format.space_before = Pt(18)
    h7.paragraph_format.space_after = Pt(6)
    r = h7.add_run("7. Research Proposals")
    r.font.size = Pt(18)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    
    add_styled_paragraph(doc, 
        "Based on the gaps identified in classical baselines and the literature review, we outline three concrete, "
        "falsifiable research proposals. One of these will be selected for detailed execution in the next phase."
    )
    
    # Proposal 1
    add_styled_paragraph(doc, "Research Proposal 1: Sequence Models for Dynamic Time-Series Momentum in Crypto Spot", bold=True)
    add_styled_paragraph(doc, 
        "1. Research Question: Does an LSTM model with a self-attention mechanism, trained on multi-scale historical returns, "
        "generate a higher Net Sharpe Ratio than a 50/200 SMA trend-following baseline on Crypto Spot (BTC/ETH), net of 0.15% frictions?\n"
        "2. Hypotheses: Standard trend-following suffers from long drawdowns in sideways markets. We hypothesize that self-attention "
        "allows the network to recognize the regime transition (from trend to range-bound) and scale down exposure, reducing Max Drawdown.\n"
        "3. Scope: Universe: BTCUSDT and ETHUSDT. Timeframe: Daily (2018-2026). Input features: Lagged returns, rolling volatility, volume indicators.\n"
        "4. Success Criteria: Positive result: Net Sharpe > 1.0 (baseline 0.73), and Max Drawdown reduced by at least 25% (i.e. above -50%). "
        "Clean negative result: The LSTM fails to cover the transaction fee drag, resulting in a Net Sharpe < 0.73, or suffers from overfitting."
    )
    
    # Proposal 2
    add_styled_paragraph(doc, "Research Proposal 2: Gradient-Boosted Trees for Cross-Sectional Alpha in Indian Large-Caps", bold=True)
    add_styled_paragraph(doc, 
        "1. Research Question: Does a LightGBM model utilizing rolling technical features and macroeconomic covariates "
        "generate statistically significant positive excess returns (alpha) in a market-neutral Indian stock universe (TCS/INFY), "
        "after accounting for dynamic Indian tax/brokerage frictions?\n"
        "2. Hypotheses: Cointegration in pairs trading often breaks down or suffers from high turnover drag. We hypothesize that GBDTs "
        "can map non-linear relationships (e.g., order imbalance, sectoral strength, and volatility spreads) to predict daily stock returns, "
        "leading to a market-neutral portfolio with a Sharpe Ratio > 1.2.\n"
        "3. Scope: Universe: NSE large-cap IT sector (TCS, INFY, WIPRO, HCLTECH). Timeframe: Daily (2018-2026). Out of scope: intraday LOB order-flow.\n"
        "4. Success Criteria: Positive: Annualized return net of taxes > 12% with a Sharpe > 1.2. Negative: Net return is negative or underperforms "
        "a simple equal-weighted buy-and-hold sector ETF."
    )
    
    # Proposal 3
    add_styled_paragraph(doc, "Research Proposal 3: Deep Reinforcement Learning for Execution and Dynamic Capital Allocation", bold=True)
    add_styled_paragraph(doc, 
        "1. Research Question: Can a Deep Q-Network (DQN) agent, operating in an actor-critic framework, optimize execution order entries "
        "and size constraints to outperform a simple time-weighted average price (TWAP) execution baseline in crypto spot markets?\n"
        "2. Hypotheses: A reinforcement learning agent can learn to place orders limit/market adaptively based on bid-ask spreads and order book depth, "
        "reducing transaction costs and execution slippage by at least 30%.\n"
        "3. Scope: Universe: BTCUSDT. Timeframe: Intraday (1-minute bars). Input: OHLCV, Order book imbalance, trade size.\n"
        "4. Success Criteria: Positive: Execution slippage reduced by >=30% compared to TWAP, and Sharpe ratio of active trading strategies "
        "improves due to execution cost savings. Negative: Agent fails to converge or exhibits unstable policy behavior, resulting in higher execution costs than TWAP."
    )
    
    # ----------------------------------------------------
    # SECTION 8: APPENDIX & CODE REPRODUCIBILITY
    # ----------------------------------------------------
    h8 = doc.add_paragraph()
    h8.paragraph_format.space_before = Pt(18)
    h8.paragraph_format.space_after = Pt(6)
    r = h8.add_run("8. Appendix & Code Reproducibility")
    r.font.size = Pt(18)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    
    add_styled_paragraph(doc, 
        "The codebase is structured to facilitate complete reproducibility by the project mentor or professor. The file structure is as follows:"
    )
    
    add_styled_paragraph(doc, 
        "crypto_backtester/\n"
        "  ├── data_engine.py          # Data downloading, cleaning, and EDA runner\n"
        "  ├── custom_backtester.py    # From-scratch, vector-based, cost/slippage-aware simulator\n"
        "  ├── run_baselines.py        # Runs SMA Crossover, Bollinger Bands, and Pairs Trading, and outputs plots\n"
        "  ├── generate_report.py      # Programmatic generator for this Microsoft Word report\n"
        "  ├── README.md               # User guide for running the pipeline\n"
        "  ├── data/                   # Directory containing cached daily CSV data\n"
        "  └── plots/                  # Directory containing generated performance plots", italic=True
    )
    
    # Save document
    doc.save("Task2_Research_Foundation.docx")
    print("[+] Word document generated successfully at: ./Task2_Research_Foundation.docx")
    
    # Copy to Artifact directory
    try:
        import shutil
        dest = os.path.join(ARTIFACT_DIR, "Task2_Research_Foundation.docx")
        shutil.copy("Task2_Research_Foundation.docx", dest)
        print(f"[+] Word document copied to artifact folder: {dest}")
    except Exception as e:
        print(f"[!] Warning: Could not copy word document to artifact folder: {e}")

if __name__ == "__main__":
    create_report()
